"""
Resume parser: extracts text and structured data from PDF/DOCX files.
"""
from __future__ import annotations

import re
from pathlib import Path
from loguru import logger

from app.utils.helpers import (
    clean_text,
    extract_emails,
    extract_phones,
    extract_linkedin,
    extract_github,
    extract_sections,
    extract_years_of_experience,
)
from app.services.nlp_engine import (
    get_spacy,
    extract_skills_from_text,
    extract_education_level,
)


def parse_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        pages_text = []
        for page in doc:
            text = page.get_text("text")
            pages_text.append(text)
        doc.close()
        return "\n".join(pages_text)
    except Exception as e:
        logger.error(f"PDF parse error ({file_path}): {e}")
        return ""


def parse_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX parse error ({file_path}): {e}")
        return ""


def extract_text(file_path: Path, file_type: str) -> str:
    """Route to the appropriate parser."""
    if file_type == "pdf":
        text = parse_pdf(file_path)
    elif file_type == "docx":
        text = parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    return clean_text(text)


def extract_name(text: str) -> str | None:
    """Extract candidate name using spaCy NER (with heuristic fallback)."""
    nlp = get_spacy()
    first_500 = text[:500]

    if nlp is not None:
        doc = nlp(first_500)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                if 2 <= len(name.split()) <= 4:
                    return name

    # Heuristic fallback: first short all-caps-initial line without digits/email
    lines = [ln.strip() for ln in first_500.split("\n") if ln.strip()]
    for line in lines[:5]:
        words = line.split()
        if (
            2 <= len(words) <= 4
            and all(w[0].isupper() for w in words if w)
            and not any(c.isdigit() for c in line)
            and "@" not in line
        ):
            return line
    return None


def extract_education(text: str) -> list[dict]:
    """Extract education entries from resume text."""
    sections = extract_sections(text)
    edu_text = sections.get("education", "")

    degree_patterns = [
        r"(?:Bachelor|B\.?S\.?|B\.?E\.?|B\.?Tech|B\.?Sc|B\.?A\.?)[^\n]*",
        r"(?:Master|M\.?S\.?|M\.?E\.?|M\.?Tech|M\.?Sc|M\.?B\.?A\.?)[^\n]*",
        r"(?:Ph\.?D|Doctorate|Doctor)[^\n]*",
        r"(?:Associate|Diploma)[^\n]*",
    ]

    education_entries = []
    source = edu_text if edu_text else text

    for pattern in degree_patterns:
        matches = re.findall(pattern, source, re.IGNORECASE)
        for match in matches:
            entry = match.strip()
            if entry:
                year_match = re.search(r"\b(19|20)\d{2}\b", entry)
                education_entries.append({
                    "degree": entry[:200],
                    "year": year_match.group(0) if year_match else None,
                })

    return education_entries


def extract_experience_entries(text: str) -> list[dict]:
    """Extract work experience entries."""
    sections = extract_sections(text)
    exp_text = sections.get("experience", "")

    date_range_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4})\s*[-–—]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|present|current|now)"

    entries = []
    matches = list(re.finditer(date_range_pattern, exp_text or text, re.IGNORECASE))

    for match in matches:
        start_ctx = max(0, match.start() - 200)
        snippet = (exp_text or text)[start_ctx:match.end() + 200]
        entries.append({
            "start": match.group(1),
            "end": match.group(2),
            "snippet": snippet.strip()[:300],
        })

    return entries


def extract_projects(text: str) -> list[str]:
    """Extract project names/descriptions."""
    sections = extract_sections(text)
    proj_text = sections.get("projects", "")
    if not proj_text:
        return []

    lines = [l.strip() for l in proj_text.split("\n") if l.strip()]
    projects = []
    for line in lines[:10]:
        if len(line) > 10:
            projects.append(line[:200])
    return projects


def extract_certifications(text: str) -> list[str]:
    """Extract certifications."""
    sections = extract_sections(text)
    cert_text = sections.get("certifications", "")

    cert_patterns = [
        r"(?:AWS|Google|Microsoft|Azure|Oracle|Cisco|CompTIA|PMI|Scrum|ITIL)[^\n]{0,100}(?:Certified|Certificate|Certification)[^\n]*",
        r"(?:Certified|Certificate|Certification)\s+[A-Z][^\n]{0,100}",
    ]

    certs: set[str] = set()
    source = cert_text if cert_text else text
    for pattern in cert_patterns:
        for match in re.findall(pattern, source, re.IGNORECASE):
            certs.add(match.strip()[:200])

    return sorted(certs)


def parse_resume(file_path: Path, file_type: str) -> dict:
    """
    Full resume parse: extract text + structured data.
    Returns a dict with all extracted fields.
    """
    raw_text = extract_text(file_path, file_type)

    if not raw_text.strip():
        logger.warning(f"Empty text extracted from {file_path}")
        return {"raw_text": "", "error": "Could not extract text from file"}

    name = extract_name(raw_text)
    emails = extract_emails(raw_text)
    phones = extract_phones(raw_text)
    linkedin = extract_linkedin(raw_text)
    github = extract_github(raw_text)
    skills = extract_skills_from_text(raw_text)
    education = extract_education(raw_text)
    education_level = extract_education_level(raw_text)
    experience_entries = extract_experience_entries(raw_text)
    years_of_experience = extract_years_of_experience(raw_text)
    projects = extract_projects(raw_text)
    certifications = extract_certifications(raw_text)
    sections = extract_sections(raw_text)

    return {
        "raw_text": raw_text,
        "name": name,
        "email": emails[0] if emails else None,
        "phone": phones[0] if phones else None,
        "linkedin": linkedin,
        "github": github,
        "skills": skills,
        "education": education,
        "education_level": education_level,
        "experience_entries": experience_entries,
        "years_of_experience": years_of_experience,
        "projects": projects,
        "certifications": certifications,
        "sections_detected": list(sections.keys()),
        "char_count": len(raw_text),
        "word_count": len(raw_text.split()),
    }
