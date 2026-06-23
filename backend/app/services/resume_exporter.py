"""
Resume Exporter — Feature 1.
Exports an optimized resume dict to DOCX or PDF bytes.
"""
from __future__ import annotations

import io
from loguru import logger


def _contact_line(contact: dict) -> str:
    parts = [v for v in [
        contact.get("email"), contact.get("phone"),
        contact.get("linkedin"), contact.get("github"),
    ] if v]
    return " | ".join(parts)


def export_to_docx(optimized: dict) -> bytes:
    """Generate DOCX bytes from an optimized resume dict."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    name = optimized.get("name") or "Candidate"
    h = doc.add_heading(name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_str = _contact_line(optimized.get("contact", {}))
    if contact_str:
        p = doc.add_paragraph(contact_str)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_order = optimized.get("section_order") or ["skills", "experience", "projects", "education", "certifications"]

    for sec in section_order:
        if sec == "skills":
            all_skills = optimized.get("skills", {}).get("all", [])
            if all_skills:
                doc.add_heading("Skills", level=2)
                primary = optimized["skills"].get("primary", [])
                secondary = optimized["skills"].get("secondary", [])
                if primary:
                    p = doc.add_paragraph()
                    p.add_run("Key Skills: ").bold = True
                    p.add_run(", ".join(primary))
                if secondary:
                    p = doc.add_paragraph()
                    p.add_run("Additional: ").bold = True
                    p.add_run(", ".join(secondary))

        elif sec == "projects":
            projects = optimized.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=2)
                for proj in projects:
                    text = proj.get("optimized") or proj.get("original", "")
                    if text:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(text)
                        if proj.get("safety") == "REQUIRES_VERIFICATION":
                            run = p.add_run("  ⚠ VERIFY")
                            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
                            run.font.size = Pt(8)

        elif sec == "experience":
            experience = optimized.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=2)
                for exp in experience:
                    start = exp.get("start", "")
                    end = exp.get("end", "")
                    if start:
                        p = doc.add_paragraph()
                        p.add_run(f"{start} – {end}").bold = True
                    text = exp.get("optimized") or exp.get("original", "")
                    if text:
                        p2 = doc.add_paragraph(style="List Bullet")
                        p2.add_run(text)
                        if exp.get("safety") == "REQUIRES_VERIFICATION":
                            run = p2.add_run("  ⚠ VERIFY")
                            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
                            run.font.size = Pt(8)

        elif sec == "education":
            education = optimized.get("education", [])
            if education:
                doc.add_heading("Education", level=2)
                for edu in education:
                    degree = edu.get("degree", "")
                    year = edu.get("year", "")
                    text = f"{degree} ({year})" if year else degree
                    if text:
                        doc.add_paragraph(text, style="List Bullet")

        elif sec == "certifications":
            certs = optimized.get("certifications", [])
            if certs:
                doc.add_heading("Certifications", level=2)
                for cert in certs:
                    doc.add_paragraph(cert, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_to_pdf(optimized: dict) -> bytes:
    """Generate PDF bytes from an optimized resume dict."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=18,
                                textColor=HexColor("#1a1a1a"), spaceAfter=4, alignment=TA_CENTER)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"], fontSize=9,
                                   textColor=HexColor("#555555"), alignment=TA_CENTER, spaceAfter=10)
    section_style = ParagraphStyle("Section", parent=styles["Heading2"], fontSize=12,
                                   textColor=HexColor("#2563eb"), spaceBefore=10, spaceAfter=3)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10,
                                  spaceAfter=3, leftIndent=12, leading=14)
    verify_style = ParagraphStyle("Verify", parent=bullet_style, textColor=HexColor("#cc4400"))

    story = []
    story.append(Paragraph(optimized.get("name") or "Candidate", name_style))

    contact_str = _contact_line(optimized.get("contact", {}))
    if contact_str:
        story.append(Paragraph(contact_str, contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#2563eb"), spaceAfter=6))

    section_order = optimized.get("section_order") or ["skills", "experience", "projects", "education", "certifications"]

    for sec in section_order:
        if sec == "skills":
            all_skills = optimized.get("skills", {}).get("all", [])
            if all_skills:
                story.append(Paragraph("SKILLS", section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd"), spaceAfter=3))
                primary = optimized["skills"].get("primary", [])
                secondary = optimized["skills"].get("secondary", [])
                if primary:
                    story.append(Paragraph(f"<b>Key:</b> {', '.join(primary)}", body_style))
                if secondary:
                    story.append(Paragraph(f"<b>Additional:</b> {', '.join(secondary)}", body_style))

        elif sec == "projects":
            projects = optimized.get("projects", [])
            if projects:
                story.append(Paragraph("PROJECTS", section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd"), spaceAfter=3))
                for proj in projects:
                    text = proj.get("optimized") or proj.get("original", "")
                    if text:
                        needs_verify = proj.get("safety") == "REQUIRES_VERIFICATION"
                        sty = verify_style if needs_verify else bullet_style
                        suffix = " [VERIFY]" if needs_verify else ""
                        story.append(Paragraph(f"• {text}{suffix}", sty))

        elif sec == "experience":
            experience = optimized.get("experience", [])
            if experience:
                story.append(Paragraph("EXPERIENCE", section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd"), spaceAfter=3))
                for exp in experience:
                    start = exp.get("start", "")
                    end = exp.get("end", "")
                    if start:
                        story.append(Paragraph(f"<b>{start} – {end}</b>", body_style))
                    text = exp.get("optimized") or exp.get("original", "")
                    if text:
                        needs_verify = exp.get("safety") == "REQUIRES_VERIFICATION"
                        sty = verify_style if needs_verify else bullet_style
                        suffix = " [VERIFY]" if needs_verify else ""
                        story.append(Paragraph(f"• {text}{suffix}", sty))

        elif sec == "education":
            education = optimized.get("education", [])
            if education:
                story.append(Paragraph("EDUCATION", section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd"), spaceAfter=3))
                for edu in education:
                    degree = edu.get("degree", "")
                    year = edu.get("year", "")
                    text = f"{degree} ({year})" if year else degree
                    if text:
                        story.append(Paragraph(f"• {text}", bullet_style))

        elif sec == "certifications":
            certs = optimized.get("certifications", [])
            if certs:
                story.append(Paragraph("CERTIFICATIONS", section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#dddddd"), spaceAfter=3))
                for cert in certs:
                    story.append(Paragraph(f"• {cert}", bullet_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()
